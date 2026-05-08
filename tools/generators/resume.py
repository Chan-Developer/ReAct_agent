# -*- coding: utf-8 -*-
"""简历文档生成工具 v2.0。

支持生成 Word 格式的专业简历文档。
功能特点：
- LLM 内容优化
- 技能进度条
- 专业模板样式
- 自动生成个人简介

需要安装依赖: pip install python-docx
"""
from __future__ import annotations

import os
import json
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Protocol, Type, TYPE_CHECKING

from ..base import BaseTool

if TYPE_CHECKING:
    from agents import ResumeAgentOrchestrator

__all__ = ["ResumeGenerator"]


# =============================================================================
# 数据模型 (增强版)
# =============================================================================

@dataclass
class Education:
    """教育经历"""
    school: str = ""
    degree: str = ""
    major: str = ""
    start_date: str = ""
    end_date: str = ""
    gpa: str = ""  # 新增：GPA
    courses: List[str] = field(default_factory=list)  # 新增：主修课程


@dataclass
class Experience:
    """工作/实习经历"""
    company: str = ""
    position: str = ""
    start_date: str = ""
    end_date: str = ""
    description: str = ""
    highlights: List[str] = field(default_factory=list)  # 新增：工作亮点


@dataclass
class Project:
    """项目经历"""
    name: str = ""
    role: str = ""
    start_date: str = ""  # 新增
    end_date: str = ""    # 新增
    description: str = ""
    highlights: List[str] = field(default_factory=list)
    tech_stack: List[str] = field(default_factory=list)  # 新增：技术栈
    target_fit_summary: str = ""
    relevance_score: float = 0.0


@dataclass
class ResumeData:
    """简历数据模型 (增强版)"""
    # 基本信息
    name: str = ""
    phone: str = ""
    email: str = ""
    location: str = ""
    
    # 社交链接 (新增)
    linkedin: str = ""
    github: str = ""
    website: str = ""
    
    # 核心内容
    summary: str = ""
    education: List[Education] = field(default_factory=list)
    experience: List[Experience] = field(default_factory=list)
    projects: List[Project] = field(default_factory=list)
    
    # 技能
    skills: List[str] = field(default_factory=list)
    
    # 附加信息 (新增)
    certificates: List[str] = field(default_factory=list)  # 证书
    awards: List[str] = field(default_factory=list)        # 荣誉奖项
    languages: List[str] = field(default_factory=list)     # 语言能力
    interests: List[str] = field(default_factory=list)     # 兴趣爱好

    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> "ResumeData":
        """从字典创建 ResumeData 实例"""
        from resume_copilot.product import sort_awards_by_importance

        # 解析教育经历
        education_list = []
        for edu in data.get("education", []):
            if isinstance(edu, dict):
                education_list.append(Education(
                    school=edu.get("school", ""),
                    degree=edu.get("degree", ""),
                    major=edu.get("major", ""),
                    start_date=edu.get("start_date", ""),
                    end_date=edu.get("end_date", ""),
                    gpa=edu.get("gpa", ""),
                    courses=edu.get("courses", []),
                ))
        
        # 解析工作经历
        experience_list = []
        for exp in data.get("experience", []):
            if isinstance(exp, dict):
                experience_list.append(Experience(
                    company=exp.get("company", ""),
                    position=exp.get("position", ""),
                    start_date=exp.get("start_date", ""),
                    end_date=exp.get("end_date", ""),
                    description=exp.get("description", ""),
                    highlights=exp.get("highlights", []),
                ))
        
        # 解析项目经历
        project_list = []
        for proj in data.get("projects", []):
            if isinstance(proj, dict):
                project_list.append(Project(
                    name=proj.get("name", ""),
                    role=proj.get("role", ""),
                    start_date=proj.get("start_date", ""),
                    end_date=proj.get("end_date", ""),
                    description=proj.get("description", ""),
                    highlights=proj.get("highlights", []),
                    tech_stack=proj.get("tech_stack", []),
                    target_fit_summary=proj.get("target_fit_summary", ""),
                    relevance_score=float(proj.get("relevance_score", 0.0) or 0.0),
                ))
        
        # 解析 skills 列表（兼容字符串和 dict 格式）
        raw_skills = data.get("skills", [])
        skills_list = []
        for skill in raw_skills:
            if isinstance(skill, str):
                skills_list.append(skill)
            elif isinstance(skill, dict):
                skill_name = skill.get("name", "")
                if skill_name:
                    skills_list.append(skill_name)
        
        return cls(
            name=data.get("name", ""),
            phone=data.get("phone", ""),
            email=data.get("email", ""),
            location=data.get("location", ""),
            linkedin=data.get("linkedin", ""),
            github=data.get("github", ""),
            website=data.get("website", ""),
            summary=data.get("summary", ""),
            education=education_list,
            experience=experience_list,
            projects=project_list,
            skills=skills_list,
            certificates=data.get("certificates", []),
            awards=sort_awards_by_importance(data.get("awards", [])),
            languages=data.get("languages", []),
            interests=data.get("interests", []),
        )

    @property
    def contact_info(self) -> List[str]:
        """获取联系信息列表"""
        info = []
        if self.phone:
            info.append(f"📱 {self.phone}")
        if self.email:
            info.append(f"📧 {self.email}")
        if self.location:
            info.append(f"📍 {self.location}")
        return info
    
    @property
    def social_links(self) -> List[str]:
        """获取社交链接"""
        links = []
        if self.github:
            links.append(f"GitHub: {self.github}")
        if self.linkedin:
            links.append(f"LinkedIn: {self.linkedin}")
        if self.website:
            links.append(f"Website: {self.website}")
        return links


# =============================================================================
# 样式配置 (增强版)
# =============================================================================

@dataclass
class ColorScheme:
    """颜色方案 - 更适合高端求职简历"""
    primary: str = "#102A43"
    secondary: str = "#486581"
    accent: str = "#0F766E"
    text: str = "#1F2933"
    light: str = "#E6FFFA"
    success: str = "#0F766E"
    border: str = "#D9E2EC"
    muted: str = "#7B8794"


@dataclass
class FontConfig:
    """字体配置 - 编辑感更强的求职排版"""
    title_size: int = 20
    heading_size: int = 11
    subheading_size: int = 10
    body_size: int = 9          # 正文
    small_size: int = 8         # 小字（时间等）


@dataclass
class SpacingConfig:
    """间距配置（单位: Pt）"""
    margin: float = 0.45
    section_gap: int = 6
    item_gap: int = 2
    line_height: float = 1.08


@dataclass
class StyleConfig:
    """样式配置（由 LayoutAgent 动态决定）"""
    colors: ColorScheme = field(default_factory=ColorScheme)
    fonts: FontConfig = field(default_factory=FontConfig)
    spacing: SpacingConfig = field(default_factory=SpacingConfig)
    show_timeline: bool = False

    @classmethod
    def get_style(cls, style: str = "default") -> "StyleConfig":
        """获取基础样式配置（作为 AI 配置的基础）"""
        return cls()


# =============================================================================
# 文档生成器抽象基类
# =============================================================================

class BaseDocumentGenerator(ABC):
    """文档生成器抽象基类"""

    def __init__(self, style: StyleConfig):
        self.style = style

    @abstractmethod
    def generate(self, data: ResumeData, output_path: str) -> bool:
        """生成文档"""
        raise NotImplementedError


# =============================================================================
# Word 文档生成器 (增强版)
# =============================================================================

class DocxGenerator(BaseDocumentGenerator):
    """Word 文档生成器 (紧凑专业版)
    
    特点：
    - 紧凑单页布局
    - 专业配色层次
    - 正式符号图标
    - 技能进度条
    """

    # 章节图标映射（正式符号）
    SECTION_ICONS = {
        "个人简介": "▎",
        "教育背景": "▎",
        "工作经历": "▎",
        "实习经历": "▎",
        "项目经验": "▎",
        "专业技能": "▎",
        "证书资质": "▎",
        "荣誉奖项": "▎",
        "语言能力": "▎",
    }

    def generate(self, data: ResumeData, output_path: str) -> bool:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        
        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Inches(self.style.spacing.margin)
            section.bottom_margin = Inches(self.style.spacing.margin)
            section.left_margin = Inches(self.style.spacing.margin)
            section.right_margin = Inches(self.style.spacing.margin)
        
        # 生成各个部分
        self._add_header(doc, data)
        self._add_summary(doc, data)
        self._add_education(doc, data)
        self._add_experience(doc, data)
        self._add_projects(doc, data)
        self._add_skills(doc, data)
        self._add_certificates(doc, data)
        self._add_awards(doc, data)
        
        doc.save(output_path)
        return True

    def _hex_to_rgb(self, hex_color: str):
        """将十六进制颜色转换为 RGBColor"""
        from docx.shared import RGBColor
        hex_color = hex_color.lstrip("#")
        return RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )

    def _apply_run_style(
        self,
        run,
        *,
        size: int,
        color: str,
        bold: bool = False,
        italic: bool = False,
        font_name: str = "Aptos",
        east_asia_font: str = "微软雅黑",
    ) -> None:
        from docx.shared import Pt
        from docx.oxml.ns import qn

        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = self._hex_to_rgb(color)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)

    def _add_header(self, doc, data: ResumeData) -> None:
        """添加页眉（更具品牌感的顶部信息区）"""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(0)
        title.paragraph_format.space_after = Pt(3)
        
        run = title.add_run(data.name or "姓名")
        self._apply_run_style(
            run,
            size=self.style.fonts.title_size,
            color=self.style.colors.primary,
            bold=True,
        )
        
        contact_parts = []
        if data.phone:
            contact_parts.append(data.phone)
        if data.email:
            contact_parts.append(data.email)
        if data.location:
            contact_parts.append(data.location)
        
        if contact_parts:
            contact = doc.add_paragraph()
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact.paragraph_format.space_before = Pt(0)
            contact.paragraph_format.space_after = Pt(1.5)
            
            run = contact.add_run(" | ".join(contact_parts))
            self._apply_run_style(
                run,
                size=self.style.fonts.small_size,
                color=self.style.colors.secondary,
            )
        
        if data.social_links:
            social = doc.add_paragraph()
            social.alignment = WD_ALIGN_PARAGRAPH.CENTER
            social.paragraph_format.space_before = Pt(0)
            social.paragraph_format.space_after = Pt(2.5)
            
            run = social.add_run(" | ".join(data.social_links))
            self._apply_run_style(
                run,
                size=self.style.fonts.small_size,
                color=self.style.colors.accent,
            )
        
        self._add_horizontal_line(doc)

    def _add_horizontal_line(self, doc) -> None:
        """添加水平分隔线"""
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '10')
        bottom.set(qn('w:color'), self.style.colors.accent.lstrip('#'))
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_section_heading(self, doc, title: str) -> None:
        """添加章节标题 - 更有高级感的编辑式层级"""
        from docx.shared import Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(self.style.spacing.section_gap)
        p.paragraph_format.space_after = Pt(3)
        
        marker = p.add_run("▌ ")
        self._apply_run_style(
            marker,
            size=self.style.fonts.heading_size,
            color=self.style.colors.accent,
            bold=True,
        )
        run = p.add_run(title.upper())
        self._apply_run_style(
            run,
            size=self.style.fonts.heading_size,
            color=self.style.colors.primary,
            bold=True,
        )
        
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '3')
        bottom.set(qn('w:color'), self.style.colors.border.lstrip('#'))
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_summary(self, doc, data: ResumeData) -> None:
        """添加个人简介 - 紧凑版"""
        if not data.summary:
            return
            
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        self._add_section_heading(doc, "个人简介")
        
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = self.style.spacing.line_height
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        run = p.add_run(data.summary)
        self._apply_run_style(
            run,
            size=self.style.fonts.body_size,
            color=self.style.colors.text,
        )

    def _add_education(self, doc, data: ResumeData) -> None:
        """添加教育背景 - 紧凑版"""
        if not data.education:
            return
            
        from docx.shared import Pt, Inches
        
        self._add_section_heading(doc, "教育背景")
        
        for edu in data.education:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            
            # 学校名称（深色）
            run = p.add_run(edu.school)
            run.bold = True
            run.font.size = Pt(self.style.fonts.subheading_size)
            run.font.color.rgb = self._hex_to_rgb(self.style.colors.primary)
            
            # 时间（浅色，右侧）
            run = p.add_run("\t" * 5)
            run = p.add_run(f"{edu.start_date} - {edu.end_date}")
            run.font.size = Pt(self.style.fonts.small_size)
            run.font.color.rgb = self._hex_to_rgb(self.style.colors.secondary)
            
            # 学位 · 专业 · GPA（同一行）
            if edu.degree or edu.major:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(0)
                p2.paragraph_format.left_indent = Inches(0.1)
                
                info_parts = []
                if edu.degree:
                    info_parts.append(edu.degree)
                if edu.major:
                    info_parts.append(edu.major)
                if edu.gpa:
                    info_parts.append(f"GPA: {edu.gpa}")
                
                run = p2.add_run(" · ".join(info_parts))
                run.font.size = Pt(self.style.fonts.body_size)
                run.font.color.rgb = self._hex_to_rgb(self.style.colors.text)

    def _add_experience(self, doc, data: ResumeData) -> None:
        """添加工作经验 - 紧凑版"""
        if not data.experience:
            return
            
        from docx.shared import Pt, Inches
        
        self._add_section_heading(doc, "工作经历")
        
        for exp in data.experience:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            
            # 公司（深色加粗）
            run = p.add_run(exp.company)
            run.bold = True
            run.font.size = Pt(self.style.fonts.subheading_size)
            run.font.color.rgb = self._hex_to_rgb(self.style.colors.primary)
            
            # 职位（强调色）
            run = p.add_run(f" | {exp.position}")
            run.font.size = Pt(self.style.fonts.body_size)
            run.font.color.rgb = self._hex_to_rgb(self.style.colors.accent)
            
            # 时间（浅色）
            run = p.add_run("\t" * 4)
            run = p.add_run(f"{exp.start_date} - {exp.end_date}")
            run.font.size = Pt(self.style.fonts.small_size)
            run.font.color.rgb = self._hex_to_rgb(self.style.colors.secondary)
            
            # 工作描述
            if exp.description:
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.1)
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(0)
                
                run = p2.add_run(exp.description)
                run.font.size = Pt(self.style.fonts.body_size)
                run.font.color.rgb = self._hex_to_rgb(self.style.colors.text)
            
            # 工作亮点
            for highlight in exp.highlights:
                p3 = doc.add_paragraph()
                p3.paragraph_format.left_indent = Inches(0.15)
                p3.paragraph_format.space_before = Pt(0)
                p3.paragraph_format.space_after = Pt(0)
                
                run = p3.add_run("• ")
                run.font.size = Pt(self.style.fonts.body_size)
                run.font.color.rgb = self._hex_to_rgb(self.style.colors.accent)
                run = p3.add_run(highlight)
                run.font.size = Pt(self.style.fonts.body_size)
                run.font.color.rgb = self._hex_to_rgb(self.style.colors.text)

    def _add_projects(self, doc, data: ResumeData) -> None:
        """添加项目经验 - 以岗位相关性优先呈现"""
        if not data.projects:
            return
            
        from docx.shared import Pt, Inches
        
        self._add_section_heading(doc, "项目经验")
        
        for proj in data.projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            
            run = p.add_run(proj.name)
            self._apply_run_style(
                run,
                size=self.style.fonts.subheading_size,
                color=self.style.colors.primary,
                bold=True,
            )
            
            if proj.role:
                run = p.add_run(f" | {proj.role}")
                self._apply_run_style(
                    run,
                    size=self.style.fonts.body_size,
                    color=self.style.colors.accent,
                )
            
            if proj.start_date or proj.end_date:
                run = p.add_run("\t" * 4)
                run = p.add_run(f"{proj.start_date} - {proj.end_date}")
                self._apply_run_style(
                    run,
                    size=self.style.fonts.small_size,
                    color=self.style.colors.secondary,
                )

            if proj.target_fit_summary:
                p_fit = doc.add_paragraph()
                p_fit.paragraph_format.left_indent = Inches(0.08)
                p_fit.paragraph_format.space_before = Pt(0)
                p_fit.paragraph_format.space_after = Pt(1)
                run = p_fit.add_run("Role fit: ")
                self._apply_run_style(
                    run,
                    size=self.style.fonts.small_size,
                    color=self.style.colors.accent,
                    bold=True,
                )
                run = p_fit.add_run(proj.target_fit_summary)
                self._apply_run_style(
                    run,
                    size=self.style.fonts.small_size,
                    color=self.style.colors.secondary,
                    italic=True,
                )
            
            if proj.tech_stack:
                p_tech = doc.add_paragraph()
                p_tech.paragraph_format.left_indent = Inches(0.1)
                p_tech.paragraph_format.space_before = Pt(0)
                p_tech.paragraph_format.space_after = Pt(1)
                
                run = p_tech.add_run("技术栈: ")
                self._apply_run_style(
                    run,
                    size=self.style.fonts.small_size,
                    color=self.style.colors.secondary,
                    bold=True,
                )
                
                run = p_tech.add_run(" | ".join(proj.tech_stack))
                self._apply_run_style(
                    run,
                    size=self.style.fonts.small_size,
                    color=self.style.colors.accent,
                )
            
            if proj.description:
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.1)
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(0.5)
                
                run = p2.add_run(proj.description)
                self._apply_run_style(
                    run,
                    size=self.style.fonts.body_size,
                    color=self.style.colors.text,
                )
            
            for highlight in proj.highlights:
                p3 = doc.add_paragraph()
                p3.paragraph_format.left_indent = Inches(0.15)
                p3.paragraph_format.space_before = Pt(0)
                p3.paragraph_format.space_after = Pt(0)
                
                run = p3.add_run("◆ ")
                self._apply_run_style(
                    run,
                    size=self.style.fonts.body_size,
                    color=self.style.colors.success,
                    bold=True,
                )
                run = p3.add_run(highlight)
                self._apply_run_style(
                    run,
                    size=self.style.fonts.body_size,
                    color=self.style.colors.text,
                )

    def _add_skills(self, doc, data: ResumeData) -> None:
        """添加专业技能"""
        if not data.skills:
            return
            
        from docx.shared import Pt
        
        self._add_section_heading(doc, "专业技能")
        
        if data.skills:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            
            for i, skill in enumerate(data.skills):
                if i > 0:
                    divider = p.add_run("   ")
                    self._apply_run_style(
                        divider,
                        size=self.style.fonts.body_size,
                        color=self.style.colors.muted,
                    )
                run = p.add_run(f"[{skill}]")
                self._apply_run_style(
                    run,
                    size=self.style.fonts.body_size,
                    color=self.style.colors.accent,
                )

    def _add_certificates(self, doc, data: ResumeData) -> None:
        """添加证书资质 - 紧凑版"""
        if not data.certificates:
            return
            
        from docx.shared import Pt
        
        self._add_section_heading(doc, "证书资质")
        
        # 一行显示所有证书
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        for i, cert in enumerate(data.certificates):
            if i > 0:
                run = p.add_run(" | ")
                run.font.size = Pt(self.style.fonts.body_size)
                run.font.color.rgb = self._hex_to_rgb(self.style.colors.secondary)
            run = p.add_run(cert)
            run.font.size = Pt(self.style.fonts.body_size)
            run.font.color.rgb = self._hex_to_rgb(self.style.colors.text)

    def _add_awards(self, doc, data: ResumeData) -> None:
        """添加荣誉奖项 - 从高价值奖项到普通奖项递减展示"""
        if not data.awards:
            return
            
        from docx.shared import Pt, Inches
        
        self._add_section_heading(doc, "荣誉奖项")
        
        for index, award in enumerate(data.awards, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.05)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0.5)

            rank = p.add_run(f"{index}. ")
            self._apply_run_style(
                rank,
                size=self.style.fonts.body_size,
                color=self.style.colors.accent,
                bold=True,
            )
            run = p.add_run(award)
            self._apply_run_style(
                run,
                size=self.style.fonts.body_size,
                color=self.style.colors.text,
            )


# =============================================================================
# 生成器工厂
# =============================================================================

class DocumentGeneratorFactory:
    """文档生成器工厂"""
    
    _generators: Dict[str, Type[BaseDocumentGenerator]] = {
        "docx": DocxGenerator,
    }

    @classmethod
    def create(cls, format_type: str, style: StyleConfig) -> BaseDocumentGenerator:
        generator_class = cls._generators.get(format_type.lower())
        if not generator_class:
            raise ValueError(f"不支持的格式: {format_type}。仅支持: docx")
        return generator_class(style)


# =============================================================================
# 对外工具类 (增强版)
# =============================================================================

class ResumeGenerator(BaseTool):
    """生成 Word 格式的专业简历文档。
    
    功能特点：
    - 支持 LLM 内容优化（自动润色简历内容）
    - 多种专业模板样式
    - 技能进度条展示
    - 自动生成个人简介
    """

    RESUME_DATA_SCHEMA = """JSON格式的简历数据，包含以下字段:
{
    "name": "姓名",
    "phone": "手机号",
    "email": "邮箱",
    "location": "所在地",
    "github": "GitHub地址(可选)",
    "linkedin": "LinkedIn地址(可选)",
    "summary": "个人简介",
    "education": [
        {"school": "学校", "degree": "学位", "major": "专业", "start_date": "开始时间", "end_date": "结束时间", "gpa": "GPA(可选)"}
    ],
    "experience": [
        {"company": "公司", "position": "职位", "start_date": "开始时间", "end_date": "结束时间", "description": "工作描述", "highlights": ["亮点1"]}
    ],
    "projects": [
        {"name": "项目名", "role": "角色", "description": "项目描述", "highlights": ["亮点1"], "tech_stack": ["技术栈"]}
    ],
    "skills": ["技能1", "技能2"],
    "certificates": ["证书1"],
    "awards": ["奖项1"]
}"""

    def __init__(
        self,
        output_dir: str = "./output",
        llm: Optional[Any] = None,
        auto_optimize: bool = True,
        use_multi_agent: bool = True,
    ) -> None:
        """初始化简历生成器。
        
        Args:
            output_dir: 输出目录
            llm: LLM 实例（用于内容优化）
            auto_optimize: 是否自动优化内容
            use_multi_agent: 是否使用多Agent架构（ContentAgent + LayoutAgent）
        """
        super().__init__(
            name="generate_resume",
            description="生成 Word 格式的专业简历文档。支持模板选择和智能分页。",
            parameters={
                "type": "object",
                "properties": {
                    "resume_data": {
                        "type": "string",
                        "description": "JSON 简历数据，或使用 @layout/@optimized/@original 引用",
                    },
                    "filename": {
                        "type": "string",
                        "description": "输出文件名（不含扩展名）",
                        "default": "resume",
                    },
                    "template": {
                        "type": "string",
                        "description": "模板名称或 @selected 使用已选模板",
                        "default": "",
                    },
                    "page_preference": {
                        "type": "string",
                        "description": "页面偏好: one_page(尽量一页), two_pages, auto(自动)",
                        "enum": ["one_page", "two_pages", "auto"],
                        "default": "auto",
                    },
                    "optimize": {
                        "type": "boolean",
                        "description": "是否使用AI优化简历内容（默认开启）",
                        "default": True,
                    },
                },
                "required": ["resume_data"],
            },
        )
        self.output_dir = output_dir
        self.llm = llm
        self.auto_optimize = auto_optimize
        self.use_multi_agent = use_multi_agent
        self._orchestrator: Optional["ResumeAgentOrchestrator"] = None
        
        os.makedirs(output_dir, exist_ok=True)
        
        # 延迟初始化多Agent协调器
        if llm is not None and use_multi_agent:
            self._init_orchestrator()

    def _init_orchestrator(self) -> None:
        """初始化多Agent协调器"""
        if self.llm is None:
            return
            
        try:
            from agents import ResumeAgentOrchestrator
            self._orchestrator = ResumeAgentOrchestrator(
                llm=self.llm,
                enable_content_optimization=True,
                enable_layout_optimization=True,
            )
        except ImportError:
            pass

    def execute(
        self,
        resume_data: str,
        filename: str = "resume",
        optimize: bool = True,
        template: str = "",
        page_preference: str = "auto",
        **kwargs,  # 兼容旧参数（如 template_style）
    ) -> str:
        """生成简历文档。
        
        Args:
            resume_data: JSON 格式的简历数据，或 "@layout"/"@optimized" 引用
            filename: 输出文件名
            optimize: 是否优化内容
            template: 模板名称或 "@selected" 使用已选模板
            page_preference: 页面偏好 ("one_page", "two_pages", "auto")
            
        Returns:
            成功或失败的消息
        """
        import tempfile
        temp_dir = tempfile.gettempdir()
        
        # 1. 加载简历数据
        raw_data, error = self._load_resume_data(resume_data, temp_dir)
        if error:
            return error
        
        # 2. 提取嵌入的布局配置
        layout_config = raw_data.pop("_layout_config", None)

        from resume_copilot.product import curate_resume
        raw_data, _ = curate_resume(raw_data, "")
        
        # 3. 加载模板配置（如果指定）
        template_config = self._load_template_config(template, temp_dir)
        if template_config and not layout_config:
            layout_config = template_config
            print("[ResumeGenerator] 使用模板配置")
        
        # 4. AI 优化（如果启用且有协调器）
        optimization_result = None
        if optimize and self.auto_optimize:
            raw_data, optimization_result, orchestrator_config = self._run_optimization(raw_data)
            if orchestrator_config and not layout_config:
                layout_config = orchestrator_config
        
        # 5. 智能分页优化
        page_notes = ""
        if page_preference != "auto" or self._should_optimize_pages(raw_data, layout_config or {}):
            raw_data, layout_config, page_notes = self._optimize_for_pages(
                raw_data, layout_config or {}, page_preference
            )
        
        # 6. 创建数据模型
        try:
            data = ResumeData.from_dict(raw_data)
        except Exception as e:
            return f"❌ 数据解析失败: {type(e).__name__}: {e}"
        
        # 7. 获取样式配置
        style = StyleConfig()
        if layout_config:
            style = self._apply_layout_config(style, layout_config)
            print("[ResumeGenerator] 使用布局配置")
        else:
            print("[ResumeGenerator] 使用默认样式")
        
        # 8. 生成文档
        try:
            output_path = os.path.join(self.output_dir, f"{filename}.docx")
            generator = DocumentGeneratorFactory.create("docx", style)
            success = generator.generate(data, output_path)
            
            if success:
                abs_path = os.path.abspath(output_path)
                
                # 构建返回消息
                extra_info = []
                
                if optimization_result and optimization_result.success:
                    mode = "多Agent" if self._orchestrator else "AI"
                    extra_info.append(f"已{mode}优化")
                
                if template_config:
                    extra_info.append("已应用模板")
                
                if page_notes:
                    extra_info.append(page_notes)
                
                extra_str = f" ({', '.join(extra_info)})" if extra_info else ""
                
                suggestion_text = ""
                if optimization_result and optimization_result.success:
                    suggestions = optimization_result.content_suggestions + optimization_result.layout_suggestions
                    if suggestions:
                        suggestion_text = "\n💡 优化建议:\n" + "\n".join(f"  • {s}" for s in suggestions[:3])
                
                return (
                    f"✅ 简历已成功生成{extra_str}!\n"
                    f"📄 文件路径: {abs_path}\n"
                    f"📋 格式: DOCX"
                    f"{suggestion_text}"
                )
            return "❌ 文档生成失败"
            
        except ImportError as e:
            missing_pkg = str(e).split("'")[-2] if "'" in str(e) else "相关包"
            return f"❌ 缺少依赖包: {missing_pkg}\n请运行: pip install python-docx"
        except Exception as e:
            return f"❌ 生成文档时出错: {type(e).__name__}: {e}"
    
    def _load_resume_data(self, resume_data: str, temp_dir: str) -> tuple:
        """加载简历数据"""
        ref = resume_data.strip() if isinstance(resume_data, str) else ""
        
        ref_map = {
            "@layout": ("layout_resume.json", "布局设计后的"),
            "@optimized": ("optimized_resume.json", "优化后的"),
            "@original": ("original_resume.json", "原始"),
        }
        
        if ref in ref_map:
            filename, desc = ref_map[ref]
            filepath = os.path.join(temp_dir, filename)
            if os.path.exists(filepath):
                with open(filepath, 'r', encoding='utf-8') as f:
                    print(f"[ResumeGenerator] 使用{desc}数据")
                    return json.load(f), None
            else:
                return None, f"❌ 未找到{desc}数据"
        
        try:
            return json.loads(resume_data) if isinstance(resume_data, str) else resume_data, None
        except json.JSONDecodeError as e:
            return None, f"❌ JSON 解析失败: {e}. 提示：可以使用 \"@layout\" 引用布局后的数据。"
    
    def _load_template_config(self, template: str, temp_dir: str) -> Optional[Dict[str, Any]]:
        """加载模板配置"""
        if not template:
            return None
        
        if template.strip() == "@selected":
            layout_file = os.path.join(temp_dir, "template_layout.json")
            if os.path.exists(layout_file):
                with open(layout_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        else:
            try:
                from tools.templates import get_registry
                registry = get_registry()
                config = registry.get(template)
                if config:
                    return config.to_layout_config()
            except Exception as e:
                print(f"[ResumeGenerator] 加载模板失败: {e}")
        
        return None
    
    def _should_optimize_pages(self, raw_data: Dict, layout_config: Dict) -> bool:
        """判断是否需要分页优化"""
        try:
            from .pagination import ContentEstimator
            estimator = ContentEstimator()
            pages = estimator.estimate_pages(raw_data, layout_config)
            return pages > 1.1
        except Exception:
            return False
    
    def _optimize_for_pages(
        self,
        raw_data: Dict,
        layout_config: Dict,
        page_preference: str,
    ) -> tuple:
        """执行分页优化"""
        try:
            from .pagination import LayoutOptimizer
            optimizer = LayoutOptimizer()
            data, style, notes = optimizer.optimize_for_pages(
                raw_data, layout_config, target=page_preference
            )
            return data, style, notes
        except Exception as e:
            print(f"[ResumeGenerator] 分页优化失败: {e}")
            return raw_data, layout_config, ""
    
    def _run_optimization(self, raw_data: Dict[str, Any]) -> tuple:
        """运行多Agent优化流程
        
        Returns:
            (优化后的数据, 优化结果对象, 布局配置)
        """
        optimization_result = None
        layout_config = None
        
        if self._orchestrator:
            try:
                result = self._orchestrator.optimize(raw_data)
                if result.success:
                    raw_data = result.optimized_resume
                    layout_config = result.layout_config
                    optimization_result = result
                    print(f"[ResumeGenerator] 多Agent优化完成，耗时 {result.execution_time:.2f}s")
            except Exception as e:
                print(f"[ResumeGenerator] 多Agent优化失败: {e}")
        
        return raw_data, optimization_result, layout_config
    
    def _apply_layout_config(self, style: StyleConfig, layout_config: Dict[str, Any]) -> StyleConfig:
        """应用 LayoutAgent 的布局配置到样式
        
        Args:
            style: 基础样式配置
            layout_config: LayoutAgent 生成的配置，包含：
                - font_config: 字体配置
                - spacing_config: 间距配置
                - visual_elements: 视觉元素开关
        """
        try:
            color_scheme = layout_config.get("color_scheme", "")
            if color_scheme:
                if color_scheme in {"executive", "professional"}:
                    style.colors = ColorScheme()
                elif color_scheme == "monochrome":
                    style.colors = ColorScheme(
                        primary="#111827",
                        secondary="#4B5563",
                        accent="#111827",
                        text="#111827",
                        light="#F3F4F6",
                        success="#111827",
                        border="#D1D5DB",
                        muted="#6B7280",
                    )
                elif color_scheme == "vibrant":
                    style.colors = ColorScheme(
                        primary="#0B1F3A",
                        secondary="#475569",
                        accent="#C2410C",
                        text="#1E293B",
                        light="#FFF7ED",
                        success="#0F766E",
                        border="#E2E8F0",
                        muted="#64748B",
                    )

            # 应用字体配置
            if "font_config" in layout_config:
                font_cfg = layout_config["font_config"]
                if "title_size" in font_cfg:
                    style.fonts.title_size = font_cfg["title_size"]
                if "heading_size" in font_cfg:
                    style.fonts.heading_size = font_cfg["heading_size"]
                if "body_size" in font_cfg:
                    style.fonts.body_size = font_cfg["body_size"]
                if "subheading_size" in font_cfg:
                    style.fonts.subheading_size = font_cfg["subheading_size"]
                if "small_size" in font_cfg:
                    style.fonts.small_size = font_cfg["small_size"]
            
            # 应用间距配置
            if "spacing_config" in layout_config:
                spacing_cfg = layout_config["spacing_config"]
                if "margin" in spacing_cfg:
                    style.spacing.margin = spacing_cfg["margin"]
                if "section_gap" in spacing_cfg:
                    style.spacing.section_gap = spacing_cfg["section_gap"]
                if "item_gap" in spacing_cfg:
                    style.spacing.item_gap = spacing_cfg["item_gap"]
                if "line_height" in spacing_cfg:
                    style.spacing.line_height = spacing_cfg["line_height"]
            
            # 应用视觉元素配置
            if "visual_elements" in layout_config:
                visual_cfg = layout_config["visual_elements"]
                if "use_timeline" in visual_cfg:
                    style.show_timeline = visual_cfg["use_timeline"]
                    
        except Exception as e:
            print(f"[ResumeGenerator] 应用布局配置失败: {e}")
        
        return style
