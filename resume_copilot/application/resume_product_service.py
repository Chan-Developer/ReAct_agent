# -*- coding: utf-8 -*-
"""Product application service for Resume Copilot."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from workflows import ResumePipeline
from workflows.base import WorkflowResult

from .reporting import build_resume_diagnosis_report
from resume_copilot.domain import normalize_resume_data
from resume_copilot.quality import ResumeBenchmarkRunner


class ResumeProductService:
    """Single product service around the core resume pipeline."""

    def __init__(self, llm=None, output_dir: str = "./storage/exports"):
        self.llm = llm
        self.output_dir = output_dir

    def generate_resume(
        self,
        resume_data: dict[str, Any],
        job_description: str = "",
        template_name: str = "",
        page_preference: str = "auto",
        output_format: str = "docx",
    ):
        normalized = normalize_resume_data(resume_data)
        if output_format.lower() in {"latex", "tex"}:
            return self._generate_latex_resume(
                normalized,
                job_description=job_description,
                template_name=template_name,
            )
        pipeline = ResumePipeline(llm=self.llm, output_dir=self.output_dir)
        result = pipeline.run(
            input_data=normalized,
            job_description=job_description,
            template_name=template_name,
            page_preference=page_preference,
            output_dir=self.output_dir,
        )
        if result.success:
            return result
        return self._generate_quick_docx(normalized, job_description, result.error)

    def _generate_latex_resume(
        self,
        resume_data: dict[str, Any],
        *,
        job_description: str = "",
        template_name: str = "",
    ) -> WorkflowResult:
        """Generate an editable LaTeX resume source file."""
        output_dir = Path(self.output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        name = str(resume_data.get("name") or "candidate").strip() or "candidate"
        path = output_dir / f"{self._safe_filename(name)}_{template_name or 'resume'}.tex"

        content = self._build_latex_document(
            resume_data,
            job_description=job_description,
            template_name=template_name,
        )
        path.write_text(content, encoding="utf-8")
        return WorkflowResult(
            success=True,
            output={
                "resume_data": resume_data,
                "output_path": str(path),
                "output_format": "latex",
            },
            suggestions=[
                "Generated editable LaTeX source. Compile with XeLaTeX if you need PDF output.",
                "LaTeX is best for stable one-page resumes, version control, and precise spacing.",
            ],
            steps_completed=1,
            total_steps=1,
        )

    def _build_latex_document(
        self,
        resume_data: dict[str, Any],
        *,
        job_description: str = "",
        template_name: str = "",
    ) -> str:
        name = self._latex_escape(str(resume_data.get("name") or "Candidate"))
        email = self._latex_escape(str(resume_data.get("email") or ""))
        phone = self._latex_escape(str(resume_data.get("phone") or ""))
        summary = self._latex_escape(str(resume_data.get("summary") or ""))
        skills = resume_data.get("skills", []) or []
        section_order = self._latex_section_order(template_name)

        sections: list[str] = []
        for section in section_order:
            if section == "summary" and summary:
                sections.append(self._latex_section("Summary", [summary]))
            elif section == "education":
                sections.extend(self._latex_items_section("Education", resume_data.get("education", []) or []))
            elif section == "experience":
                sections.extend(self._latex_items_section("Experience", resume_data.get("experience", []) or []))
            elif section == "projects":
                sections.extend(self._latex_items_section("Projects", resume_data.get("projects", []) or []))
            elif section == "skills" and skills:
                sections.append(self._latex_section("Skills", [self._latex_escape(" / ".join(str(item) for item in skills[:18]))]))

        if job_description:
            sections.append(
                self._latex_section(
                    "Target Alignment",
                    ["Tailored against the target job description. Keep final bullets truthful and evidence-led."],
                )
            )

        contact = " \\quad ".join(part for part in [email, phone] if part)
        style = self._latex_style(template_name)
        return rf"""\documentclass[10pt]{{ctexart}}
\usepackage[margin={style['margin']}]{{geometry}}
\usepackage{{enumitem}}
\usepackage{{titlesec}}
\usepackage{{xcolor}}
\usepackage[hidelinks]{{hyperref}}
\pagestyle{{empty}}
\setlength{{\parindent}}{{0pt}}
\setlist[itemize]{{leftmargin=1.2em,itemsep=2pt,topsep=2pt}}
\definecolor{{accent}}{{HTML}}{{{style['accent']}}}
\titleformat{{\section}}{{\large\bfseries\color{{accent}}}}{{}}{{0pt}}{{}}[\titlerule]
\titlespacing*{{\section}}{{0pt}}{{8pt}}{{4pt}}
\begin{{document}}
{{\LARGE\bfseries {name}}}\\[-1pt]
{contact}

{chr(10).join(sections)}
\end{{document}}
"""

    def _latex_section_order(self, template_name: str) -> list[str]:
        orders = {
            "fresh_graduate": ["summary", "education", "projects", "experience", "skills"],
            "tech_modern": ["summary", "projects", "skills", "experience", "education"],
            "minimal": ["summary", "experience", "projects", "education", "skills"],
            "management": ["summary", "experience", "projects", "skills", "education"],
        }
        return orders.get(template_name or "", orders["minimal"])

    def _latex_style(self, template_name: str) -> dict[str, str]:
        styles = {
            "fresh_graduate": {"accent": "1D4ED8", "margin": "0.62in"},
            "tech_modern": {"accent": "0E5D56", "margin": "0.55in"},
            "minimal": {"accent": "334155", "margin": "0.65in"},
            "management": {"accent": "9A3412", "margin": "0.58in"},
        }
        return styles.get(template_name or "", styles["minimal"])

    def _latex_items_section(self, title: str, items: list[Any]) -> list[str]:
        if not items:
            return []
        lines = [f"\\section*{{{title}}}"]
        for item in items:
            if not isinstance(item, dict):
                lines.append(self._latex_escape(str(item)) + r"\\")
                continue
            heading_parts = [
                item.get("name"),
                item.get("company"),
                item.get("position"),
                item.get("role"),
            ]
            heading = " / ".join(str(part).strip() for part in heading_parts if str(part or "").strip())
            if heading:
                lines.append(rf"\textbf{{{self._latex_escape(heading)}}}\\")
            description = str(item.get("description") or "").strip()
            highlights = item.get("highlights", []) or []
            if description:
                lines.append(self._latex_escape(description))
            if highlights:
                lines.append(r"\begin{itemize}")
                lines.extend(rf"\item {self._latex_escape(str(highlight))}" for highlight in highlights[:4])
                lines.append(r"\end{itemize}")
        return ["\n".join(lines)]

    def _latex_section(self, title: str, paragraphs: list[str]) -> str:
        body = "\n\n".join(paragraphs)
        return f"\\section*{{{title}}}\n{body}"

    def _latex_escape(self, value: str) -> str:
        replacements = {
            "\\": r"\textbackslash{}",
            "&": r"\&",
            "%": r"\%",
            "$": r"\$",
            "#": r"\#",
            "_": r"\_",
            "{": r"\{",
            "}": r"\}",
            "~": r"\textasciitilde{}",
            "^": r"\textasciicircum{}",
        }
        return "".join(replacements.get(char, char) for char in value)

    def evaluate_resume(
        self,
        resume_data: dict[str, Any],
        job_description: str,
        role: str = "Target Role",
        market: str = "global",
    ):
        return build_resume_diagnosis_report(
            normalize_resume_data(resume_data),
            job_description,
            role=role,
            market=market,
        )

    def run_benchmark(self, dataset_path: str = "storage/benchmarks/resume_eval_set.json"):
        runner = ResumeBenchmarkRunner()
        return runner.run(dataset_path)

    def _generate_quick_docx(
        self,
        resume_data: dict[str, Any],
        job_description: str,
        pipeline_error: str | None = None,
    ) -> WorkflowResult:
        """Generate a minimal usable DOCX when the full pipeline is unavailable."""
        try:
            from docx import Document
            from docx.shared import Pt
        except Exception as exc:  # pragma: no cover - dependency-specific fallback
            return WorkflowResult(
                success=False,
                output={"resume_data": resume_data, "output_path": ""},
                suggestions=[],
                error=f"Quick DOCX fallback unavailable: {exc}",
            )

        output_dir = Path(self.output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        name = str(resume_data.get("name") or "candidate").strip() or "candidate"
        path = output_dir / f"{self._safe_filename(name)}_quick_resume.docx"

        document = Document()
        styles = document.styles
        styles["Normal"].font.name = "Microsoft YaHei"
        styles["Normal"].font.size = Pt(10.5)

        document.add_heading(name, level=0)
        if resume_data.get("email"):
            document.add_paragraph(str(resume_data["email"]))

        document.add_heading("求职摘要", level=1)
        summary = str(resume_data.get("summary") or "").strip()
        document.add_paragraph(summary or "围绕目标岗位整理项目、技能与可量化成果。")

        document.add_heading("核心经历", level=1)
        for item in resume_data.get("experience", []) or []:
            if not isinstance(item, dict):
                continue
            title = " / ".join(
                part
                for part in [
                    str(item.get("company", "")).strip(),
                    str(item.get("position", "")).strip(),
                ]
                if part
            )
            if title:
                document.add_paragraph(title, style="List Bullet")
            description = str(item.get("description", "")).strip()
            if description:
                document.add_paragraph(description)
            for highlight in item.get("highlights", [])[:5]:
                document.add_paragraph(str(highlight), style="List Bullet")

        skills = resume_data.get("skills", []) or []
        if skills:
            document.add_heading("技能关键词", level=1)
            document.add_paragraph(" / ".join(str(item) for item in skills[:16]))

        if job_description:
            document.add_heading("岗位对齐提示", level=1)
            document.add_paragraph("已根据目标 JD 生成，可继续补充系统设计、项目经验、沟通协作等证据。")

        document.save(path)
        suggestions = [
            "已使用快速 DOCX 兜底生成，建议继续补充姓名、邮箱、教育背景和项目量化结果。",
            "下一步应把每条经历改成：背景、动作、技术栈、结果。",
        ]
        if pipeline_error:
            suggestions.append(f"Full pipeline fallback reason: {pipeline_error}")
        return WorkflowResult(
            success=True,
            output={"resume_data": resume_data, "output_path": str(path)},
            suggestions=suggestions[:5],
            steps_completed=1,
            total_steps=1,
        )

    def _safe_filename(self, value: str) -> str:
        slug = re.sub(r"[^a-zA-Z0-9\u4e00-\u9fff]+", "_", value).strip("_")
        return slug or "resume"
